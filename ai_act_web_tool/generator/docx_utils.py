from __future__ import annotations

from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt


def create_document(title: str) -> Document:
    doc = Document()
    doc.add_heading(title, level=0)
    return doc


def add_metadata_table(doc: Document, answers: dict[str, str]) -> None:
    doc.add_heading("Stammdaten", level=1)
    rows = [
        ("Projektname", answers.get("company.company_name", "—")),
        ("Unternehmen", answers.get("company.company_name", "—")),
        ("KI-System", answers.get("ai_system_profile.ai_system_name", "—")),
        ("Version", answers.get("ai_system_profile.ai_system_version", "—")),
        ("Erstellungsdatum", datetime.now().strftime("%d.%m.%Y")),
        ("Dokumentationsversion", answers.get("documentation_governance.document_version", "—")),
        ("Freigabestatus", answers.get("documentation_governance.approval_status", "Entwurf")),
        ("Verantwortliche Person", answers.get("company.contact_person", "—")),
    ]
    add_key_value_table(doc, rows)


def add_section(doc: Document, title: str, body: str) -> None:
    doc.add_heading(title, level=1)
    for paragraph in body.split("\n"):
        text = paragraph.strip()
        if text:
            doc.add_paragraph(text)


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Feld"
    header_cells[1].text = "Angabe"

    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value or "—"


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        if item.strip():
            doc.add_paragraph(item.strip(), style="List Bullet")


def add_disclaimer(doc: Document) -> None:
    doc.add_heading("Disclaimer", level=1)
    doc.add_paragraph(
        "Auf Basis der eingegebenen Informationen ergibt sich eine vorläufige Einschätzung. "
        "Die Dokumentation dient der strukturierten Compliance-Vorbereitung und sollte fachlich, "
        "technisch, datenschutzrechtlich und rechtlich geprüft werden. "
        "Eine finale Konformitätsbewertung ist erst nach vollständiger Prüfung aller Anforderungen möglich."
    )


def save_document(doc: Document, output_path: str) -> None:
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
    doc.save(output_path)


def format_answer_value(value: Any) -> str:
    if value is None:
        return "—"
    if isinstance(value, list):
        cleaned = [str(item).strip() for item in value if str(item).strip()]
        return ", ".join(cleaned) if cleaned else "—"
    text = str(value).strip()
    return text if text else "—"
